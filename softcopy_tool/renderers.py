from __future__ import annotations

import os
from pathlib import Path


SUPPORTED_FORMATS = {"md", "pdf", "docx"}


class RenderDependencyError(RuntimeError):
    pass


def normalize_formats(raw: str | list[str] | tuple[str, ...] | None) -> list[str]:
    if raw is None:
        values = ["md"]
    elif isinstance(raw, str):
        values = [item.strip().lower() for item in raw.split(",") if item.strip()]
    else:
        values = [str(item).strip().lower() for item in raw if str(item).strip()]
    unknown = sorted(set(values) - SUPPORTED_FORMATS)
    if unknown:
        raise ValueError(f"Unsupported output format(s): {', '.join(unknown)}")
    return values or ["md"]


def render_optional_outputs(markdown_path: Path, formats: list[str]) -> list[str]:
    written: list[str] = []
    text = markdown_path.read_text(encoding="utf-8")
    if "pdf" in formats:
        _render_pdf(markdown_path.with_suffix(".pdf"), text)
        written.append(str(markdown_path.with_suffix(".pdf")))
    if "docx" in formats:
        _render_docx(markdown_path.with_suffix(".docx"), text)
        written.append(str(markdown_path.with_suffix(".docx")))
    return written


def _render_pdf(path: Path, text: str) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RenderDependencyError(
            "PDF rendering requires optional dependency `reportlab`. "
            "Install with `python3 -m pip install '.[render]'`."
        ) from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    font_name = _register_pdf_font()
    font_size = 10
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    x = 48
    y = height - 48
    max_width = width - (2 * x)
    pdf.setFont(font_name, font_size)
    for line in text.splitlines():
        for segment in _wrap_pdf_line(pdf, line, font_name, font_size, max_width):
            if y < 48:
                pdf.showPage()
                pdf.setFont(font_name, font_size)
                y = height - 48
            pdf.drawString(x, y, segment)
            y -= 14
    pdf.save()


def _wrap_pdf_line(pdf: object, line: str, font_name: str, font_size: int, max_width: float) -> list[str]:
    if line == "":
        return [""]
    parts: list[str] = []
    current = ""
    for char in line:
        candidate = current + char
        if current and _pdf_string_width(pdf, candidate, font_name, font_size) > max_width:
            parts.append(current)
            current = char
        else:
            current = candidate
    if current:
        parts.append(current)
    return parts


def _pdf_string_width(pdf: object, text: str, font_name: str, font_size: int) -> float:
    string_width = getattr(pdf, "stringWidth", None)
    if callable(string_width):
        return float(string_width(text, font_name, font_size))
    return len(text) * font_size * 0.5


def _register_pdf_font() -> str:
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError as exc:
        raise RenderDependencyError(
            "PDF rendering requires optional dependency `reportlab`. "
            "Install with `python3 -m pip install '.[render]'`."
        ) from exc

    configured = os.environ.get("SOFTCOPY_CJK_FONT")
    if configured:
        font_path = Path(configured).expanduser()
        if not font_path.exists():
            raise RenderDependencyError(f"Configured CJK font does not exist: {font_path}")
        try:
            pdfmetrics.registerFont(TTFont("SoftCopyCJK", str(font_path)))
            return "SoftCopyCJK"
        except Exception as exc:
            raise RenderDependencyError(f"Could not load configured CJK font: {font_path}") from exc

    failures: list[str] = []
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception as exc:
        failures.append(f"STSong-Light: {exc}")

    for font_path in _candidate_cjk_font_paths():
        if not font_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("SoftCopyCJK", str(font_path)))
            return "SoftCopyCJK"
        except Exception as exc:
            failures.append(f"{font_path}: {exc}")

    detail = f" Tried fallback fonts: {'; '.join(failures)}." if failures else ""
    raise RenderDependencyError(
        "PDF rendering needs a CJK-capable font. "
        "Set SOFTCOPY_CJK_FONT=/path/to/font.ttf."
        f"{detail}"
    )


def _candidate_cjk_font_paths() -> list[Path]:
    return [
        Path("/Library/Fonts/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttf"),
    ]


def _render_docx(path: Path, text: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise RenderDependencyError(
            "DOCX rendering requires optional dependency `python-docx`. "
            "Install with `python3 -m pip install '.[render]'`."
        ) from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_docx_page(document, text, Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH)
    in_code_block = False
    for line in text.splitlines():
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            _add_docx_paragraph(document, line, font_name="Consolas", size_pt=9, line_spacing=Pt(12), space_after=0)
        elif line.startswith("# "):
            _add_docx_heading(document, line[2:].strip(), 1)
        elif line.startswith("## "):
            _add_docx_heading(document, line[3:].strip(), 2)
        elif line.startswith("### "):
            _add_docx_heading(document, line[4:].strip(), 3)
        elif line.startswith("- [ ] "):
            _add_docx_paragraph(document, f"□ {line[6:].strip()}")
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:].strip())
            _style_docx_run(run, "SimSun", 10.5)
            _style_docx_paragraph(paragraph, line_spacing=1.25, space_after=6)
        elif line == "":
            _add_docx_paragraph(document, "")
        else:
            _add_docx_paragraph(document, line)
    document.save(path)


def _configure_docx_page(
    document: object,
    text: str,
    Cm: object,
    Pt: object,
    RGBColor: object,
    WD_ALIGN_PARAGRAPH: object,
) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    _set_east_asia_font(normal.element, "SimSun")

    header_text = _docx_header_text(text)
    header = section.header.paragraphs[0]
    header.text = header_text
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        _style_docx_run(run, "SimSun", 10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer)


def _docx_header_text(text: str) -> str:
    title = ""
    software_name = ""
    version = ""
    for line in text.splitlines():
        if not title and line.startswith("# "):
            title = line[2:].strip()
        elif line.startswith("- 软件名称："):
            software_name = line.split("：", 1)[1].strip()
        elif line.startswith("- 版本号："):
            version = line.split("：", 1)[1].strip()
    parts = [part for part in [software_name, version, title] if part and part != "待补充"]
    return " ".join(parts) if parts else title


def _add_docx_heading(document: object, text: str, level: int) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph = document.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_docx_run(run, "SimHei", 16, bold=True)
        _style_docx_paragraph(paragraph, line_spacing=1.25, space_after=12)
    elif level == 2:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _style_docx_run(run, "SimHei", 14, bold=True)
        _style_docx_paragraph(paragraph, line_spacing=1.25, space_after=8)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _style_docx_run(run, "SimHei", 12, bold=True)
        _style_docx_paragraph(paragraph, line_spacing=1.25, space_after=6)


def _add_docx_paragraph(
    document: object,
    text: str,
    *,
    font_name: str = "SimSun",
    size_pt: float = 10.5,
    line_spacing: object = 1.25,
    space_after: int = 6,
) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _style_docx_run(run, font_name, size_pt)
    _style_docx_paragraph(paragraph, line_spacing=line_spacing, space_after=space_after)


def _style_docx_paragraph(paragraph: object, *, line_spacing: object, space_after: int) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)


def _style_docx_run(run: object, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    from docx.shared import Pt, RGBColor

    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    _set_east_asia_font(run._element, font_name)


def _set_east_asia_font(element: object, font_name: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _add_page_number(paragraph: object) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run("第 ")
    _style_docx_run(run, "SimSun", 10.5)
    field_run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char_begin)
    field_run._r.append(instr_text)
    field_run._r.append(fld_char_separate)
    field_run._r.append(fld_char_end)
    _style_docx_run(field_run, "SimSun", 10.5)
    end_run = paragraph.add_run(" 页")
    _style_docx_run(end_run, "SimSun", 10.5)
